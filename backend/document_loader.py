from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


@dataclass
class LoadedDocument:
    source: str
    text: str


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    if suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    else:
        text = path.read_text(encoding="utf-8")

    return LoadedDocument(source=path.name, text=text)


def discover_documents(directory: Path) -> list[Path]:
    return [path for path in directory.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS]


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

